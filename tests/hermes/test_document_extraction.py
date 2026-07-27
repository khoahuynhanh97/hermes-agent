from __future__ import annotations

import tempfile
import unittest
from pathlib import Path


class DocumentExtractionTests(unittest.TestCase):
    @staticmethod
    def _worker():
        from core.job_watcher import JobWorker

        return JobWorker.__new__(JobWorker)

    def test_extracts_docx_paragraphs(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "lesson.docx"
            document = Document()
            document.add_heading("Hermes knowledge", level=1)
            document.add_paragraph("Approved knowledge is searched first.")
            document.save(path)

            text = self._worker()._extract_local_text_source(path)

        self.assertIn("Hermes knowledge", text)
        self.assertIn("Approved knowledge is searched first.", text)

    def test_extracts_pdf_text(self) -> None:
        from pypdf import PdfReader, PdfWriter

        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "source.pdf"
            path = Path(temp_dir) / "lesson.pdf"
            source.write_bytes(self._minimal_pdf("Hermes PDF knowledge"))
            reader = PdfReader(source)
            writer = PdfWriter()
            writer.append_pages_from_reader(reader)
            with path.open("wb") as output:
                writer.write(output)

            text = self._worker()._extract_local_text_source(path)

        self.assertIn("Hermes PDF knowledge", text)

    @staticmethod
    def _minimal_pdf(text: str) -> bytes:
        stream = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET\n".encode("ascii")
        objects = [
            b"<< /Type /Catalog /Pages 2 0 R >>",
            b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
            (
                b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
                b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>"
            ),
            b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
            b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"endstream",
        ]
        payload = bytearray(b"%PDF-1.4\n")
        offsets = [0]
        for number, body in enumerate(objects, start=1):
            offsets.append(len(payload))
            payload.extend(f"{number} 0 obj\n".encode("ascii"))
            payload.extend(body)
            payload.extend(b"\nendobj\n")
        xref = len(payload)
        payload.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
        payload.extend(b"0000000000 65535 f \n")
        for offset in offsets[1:]:
            payload.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
        payload.extend(
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n".encode("ascii")
        )
        return bytes(payload)


if __name__ == "__main__":
    unittest.main()
